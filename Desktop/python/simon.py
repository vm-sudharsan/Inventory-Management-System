from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def collect_user_data():
    data = {}
    data['name'] = input("Enter your full name: ")
    data['email'] = input("Enter your email address: ")
    data['phone'] = input("Enter your phone number: ")
    data['address'] = input("Enter your address: ")
    data['summary'] = input("Enter a brief summary about yourself: ")

    data['education'] = []
    while True:
        edu = {}
        edu['degree'] = input("Enter your degree: ")
        edu['institution'] = input("Enter the name of the institution: ")
        edu['year'] = input("Enter the year of graduation: ")
        data['education'].append(edu)
        more = input("Do you want to add more education details? (yes/no): ")
        if more.lower() != 'yes':
            break


    data['experience'] = []
    while True:
        exp = {}
        exp['job_title'] = input("Enter your job title: ")
        exp['company'] = input("Enter the company name: ")
        exp['years'] = input("Enter the duration (e.g., 2015-2019): ")
        exp['description'] = input("Enter the job description: ")
        data['experience'].append(exp)
        more = input("Do you want to add more job experiences? (yes/no): ")
        if more.lower() != 'yes':
            break

    data['skills'] = input("Enter your skills (comma-separated): ").split(',')

    return data

def generate_docx_resume(data):
    doc = Document()

    # Title
    title = doc.add_heading(level=0)
    run = title.add_run(data['name'])
    run.bold = True
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Contact Information
    contact_info = doc.add_paragraph()
    contact_info.add_run(f"Email: {data['email']} | Phone: {data['phone']} | Address: {data['address']}")

    # Summary
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(data['summary'])

    # Education
    doc.add_heading('Education', level=1)
    for edu in data['education']:
        edu_paragraph = doc.add_paragraph()
        edu_paragraph.add_run(f"{edu['degree']}, {edu['institution']} ({edu['year']})").bold = True

    # Experience
    doc.add_heading('Experience', level=1)
    for exp in data['experience']:
        exp_paragraph = doc.add_paragraph()
        exp_paragraph.add_run(f"{exp['job_title']} at {exp['company']} ({exp['years']})").bold = True
        exp_paragraph.add_run(f"\n{exp['description']}")

    # Skills
    doc.add_heading('Skills', level=1)
    skills_paragraph = doc.add_paragraph()
    skills_paragraph.add_run(", ".join(data['skills']))

    return doc

def save_docx_resume(filename, doc):
    doc.save(filename)

def main():
    print("Welcome to the Resume Builder!")
    user_data = collect_user_data()
    docx_resume = generate_docx_resume(user_data)
    save_docx_resume("resume.docx", docx_resume)
    print("Your resume has been successfully generated and saved as 'resume.docx'.")


main()
